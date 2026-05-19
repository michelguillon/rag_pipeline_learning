"""
docx_parser.py — turn a .docx into a flat list of paragraph records
===================================================================
Shared by analyse.py (profiling) and chunker.py (chunk assembly), so both
read the document through exactly the same lens. If they parsed differently,
review_chunks.py would preview something other than what ingest.py stores.

A "record" is a dict:
    text, style, size, bold, is_list, override, in_table, words, date
Empty non-list paragraphs (layout spacers) are dropped.

KEY INSIGHT — read every cell of a row, not just the first.
This CV's table is mixed: most rows are one merged cell (gridSpan=5), but the
role and education rows carry a SECOND column holding dates. An earlier version
read only cells[0] and silently dropped every date. We now iterate the raw
<w:tc> children of each row (one element per real cell — gridSpan is an
attribute, so a merged cell appears exactly once, no fragile de-duplication):
the first cell is the content, the last cell is the date column, paired with
the content paragraphs by index.
"""

from docx.oxml.ns import qn
from docx.table import _Cell


def effective_size(para):
    """Rendered font size in points, or None.

    ARCHITECTURAL DECISION: report the *rendered* size, not just the run size.
    A paragraph's size can come from a direct run override OR be inherited from
    its paragraph style (which may itself inherit from a base style). We resolve
    the whole chain so the size reflects what the reader actually sees — the
    signal that matters for detecting visual hierarchy.
    """
    # 1. a direct run-level override wins
    sizes = [r.font.size.pt for r in para.runs if r.font.size is not None]
    if sizes:
        return max(sizes)
    # 2. else walk the style -> base_style chain
    style = para.style
    seen = set()
    while style is not None and id(style) not in seen:
        seen.add(id(style))
        if style.font is not None and style.font.size is not None:
            return style.font.size.pt
        style = style.base_style
    return None


def has_numbering(para):
    """True if the paragraph carries a numPr element (bulleted/numbered list).

    KEY INSIGHT: this — not the 'List Paragraph' style — is the reliable bullet
    signal. In real documents bullets appear under several styles (here:
    'List Paragraph' AND 'Normal'); the numPr element is always present.
    """
    pPr = para._p.pPr
    return pPr is not None and pPr.find(qn("w:numPr")) is not None


def _date_cell_texts(tc):
    """Plain text of each non-empty <w:p> in a raw <w:tc> (the date column)."""
    texts = []
    for p in tc.findall(qn("w:p")):
        txt = "".join(node.text or "" for node in p.iter(qn("w:t"))).strip()
        if txt:
            texts.append(txt)
    return texts


def extract_paragraphs(doc):
    """Flatten the document into a list of paragraph records.

    Each record: text, style, size, bold, is_list, override, in_table, words,
    date. Walks document-body paragraphs and table cells. Empty non-list
    paragraphs are dropped — they are layout spacers.
    """
    records = []

    def add(para, in_table, date=""):
        text = para.text.strip()
        is_list = has_numbering(para)
        if not text and not is_list:
            return  # spacer paragraph
        records.append({
            "text": text,
            "style": para.style.name if para.style else "(none)",
            "size": effective_size(para),
            "bold": any(r.bold for r in para.runs),
            "is_list": is_list,
            "override": any(r.font.size is not None for r in para.runs),
            "in_table": in_table,
            "words": len(text.split()),
            "date": date,
        })

    for para in doc.paragraphs:
        add(para, in_table=False)

    for table in doc.tables:
        for tr in table._tbl.tr_lst:
            tcs = tr.tc_lst
            if not tcs:
                continue
            # First cell = content; last cell = the date column (when the row
            # has more than one cell). Middle cells are empty layout padding.
            content_cell = _Cell(tcs[0], table)
            dates = _date_cell_texts(tcs[-1]) if len(tcs) > 1 else []
            date_idx = 0
            for para in content_cell.paragraphs:
                text = para.text.strip()
                is_list = has_numbering(para)
                if not text and not is_list:
                    continue  # spacer
                # Dates pair, in order, with the non-bullet content paragraphs
                # (titles / body lines) — bullets live in separate one-cell rows.
                date = ""
                if not is_list and date_idx < len(dates):
                    date = dates[date_idx]
                    date_idx += 1
                add(para, in_table=True, date=date)

    return records
