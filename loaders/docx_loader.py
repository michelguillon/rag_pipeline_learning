"""
loaders/docx_loader.py — turn a .docx into a list of Paragraph objects
======================================================================
A format-specific loader: it knows everything about python-docx and nothing
about chunking. Public interface is a single function:

    load_docx(path: str) -> list[Paragraph]

This is the refactor of the Phase 1 `docx_parser.py`. The parsing logic is
UNCHANGED — same style-inheritance walk, same numPr bullet detection, same
table date-column pairing. Only the output type changed: it now emits the
common `Paragraph` model (models/paragraph.py) instead of raw dicts, so the
rest of the pipeline is format-agnostic (SPEC_PHASE2.md, Decision 2).

A Paragraph is dropped if it is an empty non-list paragraph — a layout spacer.

KEY INSIGHT — read every cell of a row, not just the first.
This CV's table is mixed: most rows are one merged cell (gridSpan=5), but the
role and education rows carry a SECOND column holding dates. An earlier version
read only cells[0] and silently dropped every date. We now iterate the raw
<w:tc> children of each row (one element per real cell — gridSpan is an
attribute, so a merged cell appears exactly once, no fragile de-duplication):
the first cell is the content, the last cell is the date column, paired with
the content paragraphs by index.
"""

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell

from models import Paragraph


def effective_size(para):
    """Rendered font size in points, or None.

    ARCHITECTURAL DECISION: report the *rendered* size, not just the run size.
    A paragraph's size can come from a direct run override OR be inherited from
    its paragraph style (which may itself inherit from a base style). We resolve
    the whole chain so the size reflects what the reader actually sees — the
    signal that matters for detecting visual hierarchy.
    """
    # 1. a direct run-level override wins
    sizes = [r.font.size.pt for r in para.runs if r.font.size is not None]
    if sizes:
        return max(sizes)
    # 2. else walk the style -> base_style chain
    style = para.style
    seen = set()
    while style is not None and id(style) not in seen:
        seen.add(id(style))
        if style.font is not None and style.font.size is not None:
            return style.font.size.pt
        style = style.base_style
    return None


def has_numbering(para):
    """True if the paragraph carries a numPr element (bulleted/numbered list).

    KEY INSIGHT: this — not the 'List Paragraph' style — is the reliable bullet
    signal. In real documents bullets appear under several styles (here:
    'List Paragraph' AND 'Normal'); the numPr element is always present.
    """
    pPr = para._p.pPr
    return pPr is not None and pPr.find(qn("w:numPr")) is not None


def _date_cell_texts(tc):
    """Plain text of each non-empty <w:p> in a raw <w:tc> (the date column)."""
    texts = []
    for p in tc.findall(qn("w:p")):
        txt = "".join(node.text or "" for node in p.iter(qn("w:t"))).strip()
        if txt:
            texts.append(txt)
    return texts


def _make_paragraph(para, in_table, date=""):
    """Build one Paragraph from a python-docx paragraph object.

    Returns None for an empty non-list paragraph (a layout spacer), so the
    caller can skip it. `style_name` keeps the Phase 1 "(none)" sentinel for a
    styleless paragraph rather than None — downstream code does
    `style.startswith("Heading")` and must not see None for a .docx.
    """
    text = para.text.strip()
    is_list = has_numbering(para)
    if not text and not is_list:
        return None  # spacer paragraph
    return Paragraph(
        text=text,
        style_name=para.style.name if para.style else "(none)",
        rendered_size=effective_size(para),
        is_bold=any(r.bold for r in para.runs),
        has_num_pr=is_list,
        in_table=in_table,
        source_format="docx",
        date=date,
        override=any(r.font.size is not None for r in para.runs),
    )


def load_docx(path):
    """Flatten a .docx file into a list of Paragraph objects.

    Walks document-body paragraphs first, then table cells. Empty non-list
    paragraphs are dropped — they are layout spacers.
    """
    doc = Document(str(path))
    paragraphs = []

    for para in doc.paragraphs:
        p = _make_paragraph(para, in_table=False)
        if p is not None:
            paragraphs.append(p)

    for table in doc.tables:
        for tr in table._tbl.tr_lst:
            tcs = tr.tc_lst
            if not tcs:
                continue
            # First cell = content; last cell = the date column (when the row
            # has more than one cell). Middle cells are empty layout padding.
            content_cell = _Cell(tcs[0], table)
            dates = _date_cell_texts(tcs[-1]) if len(tcs) > 1 else []
            date_idx = 0
            for para in content_cell.paragraphs:
                text = para.text.strip()
                is_list = has_numbering(para)
                if not text and not is_list:
                    continue  # spacer
                # Dates pair, in order, with the non-bullet content paragraphs
                # (titles / body lines) — bullets live in separate one-cell rows.
                date = ""
                if not is_list and date_idx < len(dates):
                    date = dates[date_idx]
                    date_idx += 1
                paragraphs.append(
                    _make_paragraph(para, in_table=True, date=date))

    return paragraphs
